from libraries.renderrequest import RenderRequest
from docx import Document
import os,io,base64,subprocess,shutil
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from io import BytesIO
from libraries.helper import Helper
from datetime import datetime, timezone, timedelta
import zipfile
from lxml import etree


api = RenderRequest()

class Utilities:

    @staticmethod
    def docx_to_html(plantilla_path, docx_path, valores):
        doc = Document(plantilla_path)

        def reemplazar_runs(parrafo, valores):
            for key, value in valores.items():
                placeholder = f"${{{key}}}"
                i = 0
                while i < len(parrafo.runs):
                    run = parrafo.runs[i]
                    if placeholder in run.text:
                        # Si es imagen (ej: firma en base64)
                        if key == "firma" and value:
                            run.text = run.text.replace(placeholder, "")  # limpiar placeholder
                            try:
                                image_data = base64.b64decode(value)
                                image_stream = io.BytesIO(image_data)
                                # Insertamos la imagen en el párrafo
                                run.add_picture(image_stream, width=Inches(1.5))  
                            except Exception as e:
                                print(f"Error insertando imagen: {e}")
                        else:
                            # Reemplazo de texto normal
                            run.text = run.text.replace(placeholder, str(value))
                    else:
                        # Revisar si la variable está partida en varios runs
                        combined_text = run.text
                        j = i + 1
                        while placeholder not in combined_text and j < len(parrafo.runs):
                            combined_text += parrafo.runs[j].text
                            j += 1
                        if placeholder in combined_text:
                            if key == "firma" and value:
                                for k in range(i, j):
                                    parrafo.runs[k].text = ""  # limpiar runs del placeholder
                                try:
                                    image_data = base64.b64decode(value)
                                    image_stream = io.BytesIO(image_data)
                                    parrafo.add_run().add_picture(image_stream, width=Inches(1.5))
                                except Exception as e:
                                    print(f"Error insertando imagen: {e}")
                            else:
                                new_text = combined_text.replace(placeholder, str(value))
                                index = 0
                                for k in range(i, j):
                                    run_len = len(parrafo.runs[k].text)
                                    parrafo.runs[k].text = new_text[index:index+run_len]
                                    index += run_len
                    i += 1

        # Reemplazo en párrafos
        for p in doc.paragraphs:
            reemplazar_runs(p, valores)

        # Reemplazo en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        reemplazar_runs(p, valores)

        doc.save(docx_path)

    @staticmethod
    def signature_doc(doc_path: str, signature_path: str, width_inches: float = 1.2):
        """
        Inserta la imagen signature_path en el lugar del marcador ${firma}
        dentro de doc_path (.docx). Reemplaza la primera ocurrencia encontrada
        en cada párrafo / celda. Preserva estilos en los runs.
        """
        if not os.path.isfile(doc_path):
            raise FileNotFoundError(f"Documento no encontrado: {doc_path}")
        if not os.path.isfile(signature_path):
            raise FileNotFoundError(f"Firma no encontrada: {signature_path}")

        doc = Document(doc_path)
        placeholder = "${firma}"
        width = Inches(width_inches)

        def insert_image_at_placeholder_in_paragraph(p):
            # obtener textos originales de los runs
            run_texts = [r.text for r in p.runs]
            if not run_texts:
                return False

            full_text = "".join(run_texts)
            if placeholder not in full_text:
                return False

            # posición de la primera ocurrencia
            start_index = full_text.find(placeholder)

            # nuevo texto sin el placeholder (solo la primera ocurrencia para mantener control)
            new_full_text = full_text.replace(placeholder, "", 1)

            # repartir new_full_text en los runs usando las longitudes originales
            idx = 0
            for i, orig in enumerate(run_texts):
                l = len(orig)
                # asignar parte correspondiente (si nos pasamos, asignar resto y luego '')
                segment = new_full_text[idx: idx + l] if idx < len(new_full_text) else ""
                p.runs[i].text = segment
                idx += l

            # si hay runs adicionales (poco probable), limpiarlos
            if len(p.runs) > len(run_texts):
                for r in p.runs[len(run_texts):]:
                    r.text = ""

            # calcular en qué run empieza el placeholder (usando longitudes originales)
            acc = 0
            run_start_idx = None
            for i, orig in enumerate(run_texts):
                if acc + len(orig) > start_index:
                    run_start_idx = i
                    break
                acc += len(orig)
            if run_start_idx is None:
                run_start_idx = max(0, len(p.runs) - 1)

            # crear un nuevo run con la imagen y moverlo antes del run detectado
            new_run = p.add_run()
            try:
                # add_picture sobre Run (o sobre el run nuevo)
                new_run.add_picture(signature_path, width=width)
            except Exception as e:
                # fallback: intentar con p.add_run().add_picture(...) (ambas deberían funcionar)
                try:
                    r2 = p.add_run()
                    r2.add_picture(signature_path, width=width)
                    new_run = r2
                except Exception as e2:
                    # si falla, re-lanzamos para debugging
                    raise

            # mover el run con la imagen al lugar correcto (antes del run_start)
            target_run = p.runs[run_start_idx]
            # insertar el nuevo run antes del target_run en el XML
            target_run._r.addprevious(new_run._r)

            return True

        # 1) reemplazar en párrafos del body
        for p in doc.paragraphs:
            insert_image_at_placeholder_in_paragraph(p)

        # 2) reemplazar en tablas (cada celda)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        insert_image_at_placeholder_in_paragraph(p)

        # 3) (opcional) podrías querer buscar también en headers/footers si tu plantilla lo usa

        # Guardar el documento (sobre-escribe el original)
        doc.save(doc_path)

    @staticmethod
    def docx_to_pdf_with_images(docx_path: str, pdf_path: str):
        """
        Convierte un docx a PDF preservando:
            - texto con estilos (negrita, tamaño de fuente)
            - alineación (izquierda, centro, derecha)
            - saltos de línea
            - tablas
            - imágenes con posicionamiento correcto
        """
        if not os.path.isfile(docx_path):
            raise FileNotFoundError(f"Documento no encontrado: {docx_path}")

        doc = Document(docx_path)
        c = canvas.Canvas(pdf_path, pagesize=A4)
        page_width, page_height = A4
        x_margin = 0.75 * inch
        y_position = page_height - 0.75 * inch
        
        # Extraer todas las imágenes del documento
        image_map = {}
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    image_map[rel_id] = rel.target_part.blob
        except Exception as e:
            print(f"Advertencia al extraer imágenes: {e}")

        def check_page_break(needed_space=inch):
            """Verifica si necesitamos una nueva página"""
            nonlocal y_position
            if y_position < needed_space:
                c.showPage()
                y_position = page_height - 0.75 * inch
                return True
            return False

        def get_paragraph_alignment(p):
            """Obtiene la alineación del párrafo"""
            try:
                alignment = p.alignment
                if alignment == 1:  # CENTER
                    return 'center'
                elif alignment == 2:  # RIGHT
                    return 'right'
                elif alignment == 3:  # JUSTIFY
                    return 'justify'
                else:  # LEFT o None
                    return 'left'
            except:
                return 'left'

        def draw_text_with_alignment(text, y_pos, alignment='left', font_name='Helvetica', font_size=11, bold=False):
            """Dibuja texto con alineación y estilo"""
            if bold:
                font_name = font_name + '-Bold'
            
            c.setFont(font_name, font_size)
            text_width = c.stringWidth(text, font_name, font_size)
            
            if alignment == 'center':
                x_pos = (page_width - text_width) / 2
            elif alignment == 'right':
                x_pos = page_width - x_margin - text_width
            else:  # left or justify
                x_pos = x_margin
            
            c.drawString(x_pos, y_pos, text)
            return font_size + 2  # retorna el espacio usado

        def draw_paragraph(p):
            """Dibuja un párrafo con estilos, alineación e imágenes"""
            nonlocal y_position
            
            alignment = get_paragraph_alignment(p)
            
            # Primero, procesar imágenes en el párrafo
            images_in_paragraph = []
            for run in p.runs:
                if hasattr(run, '_element'):
                    for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                        try:
                            # Obtener dimensiones de la imagen del XML
                            extent = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                            blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                            
                            if blip is not None:
                                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed_id and embed_id in image_map:
                                    # Calcular dimensiones (EMUs to inches: 914400 EMUs = 1 inch)
                                    img_width = 1.2 * inch
                                    img_height = 1.2 * inch
                                    
                                    if extent is not None:
                                        try:
                                            cx = int(extent.get('cx', 914400))
                                            cy = int(extent.get('cy', 914400))
                                            img_width = (cx / 914400.0) * inch
                                            img_height = (cy / 914400.0) * inch
                                        except:
                                            pass
                                    
                                    images_in_paragraph.append({
                                        'data': image_map[embed_id],
                                        'width': img_width,
                                        'height': img_height
                                    })
                        except Exception as e:
                            print(f"Error al procesar imagen: {e}")
            
            # Dibujar imágenes
            for img_info in images_in_paragraph:
                check_page_break(img_info['height'] + inch)
                
                img_stream = BytesIO(img_info['data'])
                img_reader = ImageReader(img_stream)
                
                # Calcular posición X según alineación
                if alignment == 'center':
                    x_pos = (page_width - img_info['width']) / 2
                elif alignment == 'right':
                    x_pos = page_width - x_margin - img_info['width']
                else:
                    x_pos = x_margin
                
                try:
                    c.drawImage(img_reader, x_pos, y_position - img_info['height'], 
                              width=img_info['width'], height=img_info['height'], 
                              preserveAspectRatio=True, mask='auto')
                    y_position -= (img_info['height'] + 5)
                except Exception as e:
                    print(f"Error dibujando imagen: {e}")
            
            # Procesar texto del párrafo
            if p.text.strip():
                # Analizar runs para detectar estilos
                text_segments = []
                for run in p.runs:
                    if run.text:
                        is_bold = run.bold if run.bold is not None else False
                        font_size = 11  # tamaño por defecto
                        
                        # Intentar obtener tamaño de fuente
                        try:
                            if run.font.size:
                                font_size = run.font.size.pt
                        except:
                            pass
                        
                        text_segments.append({
                            'text': run.text,
                            'bold': is_bold,
                            'size': font_size
                        })
                
                # Si todos los runs tienen el mismo estilo, dibujar como una unidad
                if text_segments:
                    # Combinar texto
                    full_text = ''.join([seg['text'] for seg in text_segments])
                    # Usar el estilo del primer run (simplificación)
                    is_bold = text_segments[0]['bold']
                    font_size = text_segments[0]['size']
                    
                    # Dividir en líneas
                    lines = full_text.split('\n')
                    for line in lines:
                        if line.strip():
                            # Word wrap si es necesario
                            font_name = 'Helvetica-Bold' if is_bold else 'Helvetica'
                            c.setFont(font_name, font_size)
                            
                            max_width = page_width - (2 * x_margin)
                            
                            # Dividir línea si es muy larga
                            words = line.split()
                            current_line = ""
                            
                            for word in words:
                                test_line = current_line + " " + word if current_line else word
                                test_width = c.stringWidth(test_line, font_name, font_size)
                                
                                if test_width <= max_width:
                                    current_line = test_line
                                else:
                                    if current_line:
                                        check_page_break(font_size + 10)
                                        draw_text_with_alignment(current_line, y_position, alignment, 
                                                               'Helvetica', font_size, is_bold)
                                        y_position -= (font_size + 2)
                                    current_line = word
                            
                            if current_line:
                                check_page_break(font_size + 10)
                                draw_text_with_alignment(current_line, y_position, alignment, 
                                                       'Helvetica', font_size, is_bold)
                                y_position -= (font_size + 2)
                        else:
                            # Línea vacía
                            y_position -= 7
                
                # Espacio después del párrafo
                y_position -= 6

        def draw_table(table):
            """Dibuja una tabla con bordes"""
            nonlocal y_position
            
            num_cols = len(table.columns)
            available_width = page_width - (2 * x_margin)
            col_width = available_width / num_cols if num_cols > 0 else available_width
            
            for row in table.rows:
                # Calcular contenido de celdas
                cell_contents = []
                max_lines = 1
                
                for cell in row.cells:
                    cell_text = ""
                    for p in cell.paragraphs:
                        if p.text.strip():
                            cell_text += p.text.strip() + "\n"
                    cell_text = cell_text.strip()
                    cell_contents.append(cell_text)
                    
                    # Contar líneas
                    if cell_text:
                        lines = cell_text.split('\n')
                        max_lines = max(max_lines, len(lines))
                
                # Calcular altura de la fila
                row_height = max_lines * 14 + 10
                
                # Verificar espacio
                check_page_break(row_height + inch)
                
                # Dibujar celdas
                for idx, content in enumerate(cell_contents):
                    x_pos = x_margin + (idx * col_width)
                    
                    # Dibujar borde
                    c.rect(x_pos, y_position - row_height, col_width, row_height)
                    
                    # Dibujar contenido
                    if content:
                        lines = content.split('\n')
                        temp_y = y_position - 12
                        
                        for line in lines[:max_lines]:  # Limitar líneas
                            if line.strip():
                                # Truncar texto si es muy largo
                                c.setFont('Helvetica', 10)
                                while c.stringWidth(line, 'Helvetica', 10) > (col_width - 4):
                                    line = line[:-1]
                                
                                c.drawString(x_pos + 2, temp_y, line)
                                temp_y -= 14
                
                y_position -= row_height
            
            # Espacio después de tabla
            y_position -= 10

        # Procesar documento en orden
        for element in doc.element.body:
            if element.tag.endswith('p'):
                for p in doc.paragraphs:
                    if p._element == element:
                        draw_paragraph(p)
                        break
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        draw_table(table)
                        break

        c.save()
        print(f"PDF generado correctamente: {pdf_path}")

        
    @staticmethod    
    def potectted_docx(docx_file):
        #docx_file = "entrada.docx"
        output_file = "protegido.docx"

        with zipfile.ZipFile(docx_file, 'r') as zin:
            with zipfile.ZipFile(output_file, 'w') as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)

                    # Editar document.xml
                    if item.filename == "word/settings.xml":
                        tree = etree.fromstring(data)
                        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

                        protection = etree.Element("{%s}documentProtection" % ns["w"])
                        protection.set("w:edit", "readOnly")
                        protection.set("w:enforcement", "1")

                        tree.insert(0, protection)
                        data = etree.tostring(tree)

                    zout.writestr(item, data)



    @staticmethod    
    async def formCharge(session: dict):
        info_index=[]
        schema_name = session.get("schema")
        company_id=int(session.get('company'))

        response = await api.get_data("company",id=company_id,schema="global")
        company = response['data'] if response["status"] == "success" else [] 

        getQuery=f"id={session.get('sale')}"
        response = await api.get_data("sale/informe",query=getQuery,schema=schema_name)
        venta=response['data'][0] if response["status"] == "success" else [] 

        programid=int(venta['program_id'])
        response = await api.get_data("programac",id=programid,schema=schema_name)
        program=response['data'] if response["status"] == "success" else []
        
        colegioid=int(venta['establecimiento_id'])
        response = await api.get_data("colegio",id=colegioid,schema=schema_name)
        colegio=response['data'] if response["status"] == "success" else []

        info_index = {
            "curso": f"{colegio['nombre']}-{venta['curso']}-{venta['idcurso']}",
            "programa": program["name"],
            "fechaultimo": Helper.formatear(venta["fecha_ultpag"]),
            "fechasalida": Helper.formatear(venta["fechasalida"]),
            "company": company["nomfantasia"],
            "identificador": company["identificador"]
        }

        return info_index     
    

    @staticmethod    
    async def formPaymentCharge(session: dict):
        info_index=[]
        schema_name = session.get("schema")
        company_id=int(session.get('company'))

        response = await api.get_data("company",id=company_id,schema="global")
        company = response['data'] if response["status"] == "success" else [] 

        response = await api.get_data("sale",id=int(session.get('sale')),schema=schema_name)
        venta = response['data'] if response["status"] == "success" else []

        if  session.get("position")== "General":
            consulta = int(session.get('user_curso_id'))
        else:
            consulta = int(session.get('id'))

        response = await api.get_data("curso",id=consulta,schema=schema_name)
        cursos = response['data'] if response["status"] == "success" else []

        colegioid=int(venta['establecimiento_id'])
        response = await api.get_data("colegio",id=colegioid,schema=schema_name)
        colegio=response['data'] if response["status"] == "success" else []
        
        program_id= int(venta["program_id"])
        response = await api.get_data("programac",id=program_id,schema=schema_name)
        program = response['data'] if response["status"] == "success" else []

        # Construcción de la fila
        fecha_iso = venta['fecha_ultpag']
        # Convertir y formatear
        fechaultpag = datetime.strptime(fecha_iso, "%Y-%m-%dT%H:%M:%SZ").strftime("%Y-%m-%d")

        # Construcción de la fila
        fecha_iso = venta['fechasalida']
        # Convertir y formatear
        fechasal = datetime.strptime(fecha_iso, "%Y-%m-%dT%H:%M:%SZ").strftime("%Y-%m-%d")

        info_index = {
            "curso": f"{colegio['nombre']}-{venta['curso']}-{venta['idcurso']}",
            "company": company["nomfantasia"],
            "colegio": colegio["nombre"],
            "nomcurso": f'{venta["curso"]}-{venta["idcurso"]}',
            "programa": program["name"],
            "apoderado": cursos['nombreapod'],
            "alumno": cursos['nombrealumno'],
            "rutalumno": cursos["rutalumno"], 
            "pasaporte": cursos["pasaporte"],
            "fechanac": cursos['fechanac'],
            "regionid": cursos["region_id"],
            "comunaid": cursos["comuna_id"],
            "dircalle": cursos['dircalle'],
            "dirnumero": cursos['dirnumero'],
            "nrodepto": cursos['nrodepto'],
            "fono": cursos['fono'],
            "celular": cursos['celular'],
            "correo": cursos['correo'],
            "cuotas": venta["cuotas"],
            "fechaultimo": fechaultpag,
            "fechasalida": fechasal,
            "identificador": company["identificador"]
        }
        return info_index 